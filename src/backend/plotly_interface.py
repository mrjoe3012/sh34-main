"""Here we define all plotly-related interactions such as generating charts from json input."""

#Mustafa Onur Cay - 19/10/2023
import json
import os
import shutil
import io
import plotly.graph_objs as go
import plotly.io as pio
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_plot_jsons(config_json_list,data_json):
    """
        Takes in the list of configJSONs and converts each
        of them to its respective JSON object form.
        Returns a list of JSON object representations of plots
    """

    plot_jsons = []
    for config_json in config_json_list:
        plot_jsons.append(generate_plot_json(config_json,data_json))

    return json.dumps(plot_jsons)


def generate_plot_json(config_json, data_json):
    """
        Takes in the config_json received from the frontend,
        creates the plotly figure and returns its JSON form
    """

    # Generate a Dictionary of Properties from config_json
    properties = build_property_dict(config_json)

    fig = go.Figure()

    fig = update_traces(fig, properties, data_json)
    fig = update_xaxis(fig, properties)
    fig = update_yaxis(fig, properties)
    fig = update_plot_colours(fig, properties)
    fig = update_grid_lines(fig, properties)
    fig = update_title(fig, properties)
    fig = update_xaxis_ticklabels(fig, properties)
    fig = update_yaxis_ticklabels(fig, properties)
    fig = update_plotsize(fig, properties)
    fig = update_annotations(fig, properties)

    fig_json = fig.to_json()
    return fig_json


def generate_plot_htmls(config_json_list, data_json):
    """
        Takes in the list of configJSONs and converts each of
        them to its respective HTML figure form.
        Returns a list of HTML representations of plots
    """

    plot_htmls = []
    for config_json in config_json_list:
        plot_htmls.append(generate_plot_html(config_json,data_json))

    return json.dumps(plot_htmls)


def generate_plot_html(config_json, data_json):
    """This function returns HTML form of a plotly figure"""
    fig = generate_plot(config_json, data_json)
    return fig.to_html()


def generate_plot_png(config_json, data_json):
    """This is returns png byte stream of the plotly figure"""
    fig = generate_plot(config_json, data_json)
    plot_title = build_property_dict(config_json)['plot_title']

    return pio.to_image(fig,format='png'), plot_title


def generate_plot(config_json, data_json):
    """ Takes in the config_json received from the frontend,
        creates the plotly figure and returns it
    """

    # Generate a Dictionary of Properties from config_json
    properties = build_property_dict(config_json)

    fig = go.Figure()

    fig = update_traces(fig, properties, data_json)
    fig = update_xaxis(fig, properties)
    fig = update_yaxis(fig, properties)
    fig = update_plot_colours(fig, properties)
    fig = update_grid_lines(fig, properties)
    fig = update_title(fig, properties)
    fig = update_xaxis_ticklabels(fig, properties)
    fig = update_yaxis_ticklabels(fig, properties)
    fig = update_plotsize(fig, properties)
    fig = update_annotations(fig, properties)

    return fig


def update_traces(fig, properties, data_json):
    """
        Updates the figure with the specified traces from the trace array.
        Returns the updated figure after traces have been added.
    """

    # Create the Figure, plotting the data for each trace in the trace array.

    for trace in properties["traces"]:
        plot_type = trace["plotType"]
        trace_name = trace["name"]
        trace_marker_colour = trace["markerColour"]
        plot_data_x = trace["plotDataX"]
        plot_data_y = trace["plotDataY"]

        # Generating the dataframe based on the plot_indicator field.
        # NB - still no idea what the 3rd parameter is for.

        graph_data = data_extract(data_json,plot_data_x,plot_data_y)

        x_data = graph_data["x"]
        y_data = graph_data["y"]


        # Adding the appropriate trace
        if plot_type=="Bar":
            fig.add_trace(go.Bar(x=x_data,
                                 y=y_data,
                                 name=trace_name,
                                 marker={"color": trace_marker_colour}))
        elif plot_type=="Scatter":
            fig.add_trace(go.Scatter(x=x_data,
                                     y=y_data,
                                     mode='markers',
                                     name=trace_name,
                                     marker={"color": trace_marker_colour}))
        elif plot_type=="Line":
            fig.add_trace(go.Line(x=x_data,
                                  y=y_data,
                                  name=trace_name,
                                  marker={"color": trace_marker_colour}))
        elif plot_type=="Pie":
            fig.add_trace(go.Pie(labels=x_data,
                                 values=y_data,
                                 name=trace_name,
                                 textinfo='label',
                                 hoverinfo='label+value'))
            fig.update_traces(domain={"x": [0.5,0.5], "y":[0.5,0.5]}, selector={"type": "pie"})

    return fig


def update_title(fig, properties):
    """
        Takes in the current figure, the properties dictionary and:
            - Alters the figure's title using the properties dictionary.
        Returns the updated plotly figure.
    """

    if properties["title_style_mode"] == "default":
        fig.update_layout(
            title={
                'text': properties["plot_title"],
                'font': {
                    'family': properties["title_typeface_default"],
                    'size': properties["title_size_default"],
                    'color': properties["title_colour_default"]
                }
            }
        )
    elif properties["title_style_mode"] == "custom":
        fig.update_layout(
            title={
                'text': properties["plot_title"],
                'font': {
                    'family': properties["title_typeface_custom"],
                    'size': properties["title_size_custom"],
                    'color': properties["title_colour_custom"]
                }
            }
        )

    return fig


def update_xaxis(fig, properties):
    """
        Takes in the current figure, the properties dictionary and:
            - Alters the figure's xaxis label using the properties dictionary.
        Returns the updated plotly figure.
    """

    if properties["xaxis_style_mode"] == "default":
        fig.update_xaxes(
            title_text=properties["xaxis_text"],
            title_font = {"size": properties["xaxis_size_default"],
                          "color": properties["xaxis_colour_default"],
                          "family": properties["xaxis_typeface_default"]}
        )
    elif properties["xaxis_style_mode"] == "custom":
        fig.update_xaxes(
            title_text=properties["xaxis_text"],
            title_font = {"size": properties["xaxis_size_custom"],
                          "color": properties["xaxis_colour_custom"],
                          "family": properties["xaxis_typeface_custom"]}
        )

    return fig


def update_yaxis(fig, properties):
    """
        Takes in the current figure, the properties dictionary and:
            - Alters the figure's yaxis label using the properties dictionary.
        Returns the updated plotly figure.
    """

    if properties["yaxis_style_mode"] == "default":
        fig.update_yaxes(
            title_text=properties["yaxis_text"],
            title_font = {"size": properties["yaxis_size_default"],
                          "color": properties["yaxis_colour_default"],
                          "family": properties["yaxis_typeface_default"]}
        )
    elif properties["yaxis_style_mode"] == "custom":
        fig.update_yaxes(
            title_text=properties["yaxis_text"],
            title_font = {"size": properties["yaxis_size_custom"],
                          "color": properties["yaxis_colour_custom"],
                          "family": properties["yaxis_typeface_custom"]}
        )

    return fig


def update_plotsize(fig, properties):
    """
        Takes in the current figure, the properties dictionary and:
            - Alters the figure's size using the properties dictionary.
        Returns the updated plotly figure.
    """

    fig.update_layout(width=int(properties["plot_width"]),
                      height=int(properties["plot_height"]))

    return fig


def update_plot_colours(fig, properties):
    """
        Takes in the current figure, the properties dictionary and:
            - Alters the figure's colour scheme using the properties dictionary.
        Returns the updated plotly figure.
    """

    fig.update_layout(
        plot_bgcolor = properties["plot_background_colour"],
        paper_bgcolor = properties["plot_margin_colour"]
    )

    return fig


def update_grid_lines(fig, properties):
    """
        Takes in the current figure, the properties dictionary and:
            - Alters the figure's grid lines using the properties dictionary.
        Returns the updated plotly figure.
    """

    fig.update_layout(
        xaxis = {"showgrid": properties["display_xaxis_gridlines"]},
        yaxis = {"showgrid": properties["display_yaxis_gridlines"]}
    )

    return fig


def update_xaxis_ticklabels(fig, properties):
    """
        Takes in the current figure, the properties dictionary and:
            - Alters the figure's xaxis tick labels using the properties dictionary.
        Returns the updated plotly figure.
    """

    fig.update_layout(
        xaxis = {"tickangle": properties["xaxis_ticks_angle"],
                 "side": properties["xaxis_ticks_position"]}
    )

    if properties["xaxis_ticks_style_mode"] == "default":
        fig.update_layout(
            xaxis={
                "tickfont": {
                    "family": properties["xaxis_ticks_typeface_default"],
                    "size": properties["xaxis_ticks_size_default"],
                    "color": properties["xaxis_ticks_colour_default"]
                }
            }
        )
    elif properties["xaxis_ticks_style_mode"] == "custom":
        fig.update_layout(
            xaxis={
                "tickfont": {
                    "family": properties["xaxis_ticks_typeface_custom"],
                    "size": properties["xaxis_ticks_size_custom"],
                    "color": properties["xaxis_ticks_colour_custom"]
                }
            }
        )

    return fig


def update_yaxis_ticklabels(fig, properties):
    """
        Takes in the current figure, the properties dictionary and:
            - Alters the figure's yaxis tick labels using the properties dictionary.
        Returns the updated plotly figure.
    """

    fig.update_layout(
            yaxis={
                "tickangle": properties["yaxis_ticks_angle"],
                "side": properties["yaxis_ticks_position"]
            }
    )

    if properties["yaxis_ticks_style_mode"] == "default":
        fig.update_layout(
            yaxis={
                "tickfont": {
                    "family": properties["yaxis_ticks_typeface_default"],
                    "size": properties["yaxis_ticks_size_default"],
                    "color": properties["yaxis_ticks_colour_default"]
                }
            }
        )
    elif properties["yaxis_ticks_style_mode"] == "custom":
        fig.update_layout(
            yaxis={
                "tickfont": {
                    "family": properties["yaxis_ticks_typeface_custom"],
                    "size": properties["yaxis_ticks_size_custom"],
                    "color": properties["yaxis_ticks_colour_custom"]
                }
            }
        )

    return fig


def update_annotations(fig, properties):
    """
        Creates each annotation from the annotation array in the properties dictionary
    """

    for annotation in properties["annotations"]:
        fig.add_annotation(
            x=annotation["xPos"],
            y=annotation["yPos"],
            xref=annotation["xref"],
            yref=annotation["yref"],
            text=annotation["text"],
            showarrow=annotation["showArrow"],
            arrowhead=2,
            arrowsize=1,
            arrowcolor=annotation["arrowColour"],
            ax=annotation["arrowOffsetX"],
            ay=annotation["arrowOffsetY"],
            arrowwidth=annotation["arrowWidth"],
            font={
                "color": annotation["styling"]["fontColour"],
                "size": annotation["styling"]["fontSize"],
                "family": annotation["styling"]["typeface"]
            }
        )

    return fig


def build_property_dict(config_json):
    """
        Takes in the config_json JSON, and creates a dictionary of properties from
        the config_json. This property dictionary is referenced throught methods
        in this script to access the attributes from the configJSON
    """

    properties = {}

    # Traces Array
    properties.update({"traces": config_json["traces"]})

    # Annotation Array
    properties.update({"annotations": config_json["annotations"]})

    # General Plot Options
    general_options = config_json["generalOptions"]
    properties.update({
        "plot_width": general_options["plotWidth"],
        "plot_height": general_options["plotHeight"],
        "display_xaxis_gridlines": general_options["displayXAxisLines"],
        "display_yaxis_gridlines": general_options["displayYAxisLines"],
        "xaxis_scale": general_options["xAxisScale"],
        "yaxis_scale": general_options["yAxisScale"],
    })


    # Labelling Options
    labelling_options = config_json["labellingOptions"]


    # Title
    title_options = labelling_options["title"]
    title_styling = title_options["styling"]
    properties.update({
        "plot_title": title_options["plotTitle"],
        "title_style_mode": title_styling["currentStylingMode"],
        # Default Title Styling
        "title_colour_default": title_styling["defaultFontStyle"]["fontColour"],
        "title_size_default": title_styling["defaultFontStyle"]["fontSize"],
        "title_typeface_default": title_styling["defaultFontStyle"]["typeface"],
        # Custom Title Styling
        "title_colour_custom": title_styling["customFontStyle"]["fontColour"],
        "title_size_custom": title_styling["customFontStyle"]["fontSize"],
        "title_typeface_custom": title_styling["customFontStyle"]["typeface"],
    })


    # XAxis
    xaxis_options = labelling_options["xAxis"]
    xaxis_styling = xaxis_options["styling"]
    properties.update({
        "xaxis_text": xaxis_options["xAxisText"],
        "xaxis_style_mode": xaxis_styling["currentStylingMode"],
        # Default X-Axis Styling
        "xaxis_colour_default": xaxis_styling["defaultFontStyle"]["fontColour"],
        "xaxis_size_default": xaxis_styling["defaultFontStyle"]["fontSize"],
        "xaxis_typeface_default": xaxis_styling["defaultFontStyle"]["typeface"],
        # Custom X-Axis Styling
        "xaxis_colour_custom": xaxis_styling["customFontStyle"]["fontColour"],
        "xaxis_size_custom": xaxis_styling["customFontStyle"]["fontSize"],
        "xaxis_typeface_custom": xaxis_styling["customFontStyle"]["typeface"],
    })


    # XAxis Tick Labels
    xaxis_ticks = xaxis_options["tickLabels"]
    xaxis_ticks_styling = xaxis_ticks["styling"]
    properties.update({
        "xaxis_ticks_angle": xaxis_ticks["tickAngle"],
        "xaxis_ticks_position": xaxis_ticks["tickPosition"],
        "xaxis_ticks_style_mode": xaxis_ticks_styling["currentStylingMode"],
        # Default X-Axis Tick Label Styling
        "xaxis_ticks_colour_default": xaxis_ticks_styling["customFontStyle"]["fontColour"],
        "xaxis_ticks_size_default": xaxis_ticks_styling["customFontStyle"]["fontSize"],
        "xaxis_ticks_typeface_default": xaxis_ticks_styling["customFontStyle"]["typeface"],
        # Custom X-Axis Tick Label Styling
        "xaxis_ticks_colour_custom": xaxis_ticks_styling["customFontStyle"]["fontColour"],
        "xaxis_ticks_size_custom": xaxis_ticks_styling["customFontStyle"]["fontSize"],
        "xaxis_ticks_typeface_custom": xaxis_ticks_styling["customFontStyle"]["typeface"],
    })


    # YAxis
    yaxis_options = labelling_options["yAxis"]
    yaxis_styling = yaxis_options["styling"]
    properties.update({
        "yaxis_text": yaxis_options["yAxisText"],
        "yaxis_style_mode": yaxis_styling["currentStylingMode"],
        # Default Y-Axis Styling
        "yaxis_colour_default": yaxis_styling["defaultFontStyle"]["fontColour"],
        "yaxis_size_default": yaxis_styling["defaultFontStyle"]["fontSize"],
        "yaxis_typeface_default": yaxis_styling["defaultFontStyle"]["typeface"],
        # Custom Y-Axis Styling
        "yaxis_colour_custom": yaxis_styling["customFontStyle"]["fontColour"],
        "yaxis_size_custom": yaxis_styling["customFontStyle"]["fontSize"],
        "yaxis_typeface_custom": yaxis_styling["customFontStyle"]["typeface"],
    })


    # YAxis Tick Labels
    yaxis_ticks = yaxis_options["tickLabels"]
    yaxis_ticks_styling = yaxis_ticks["styling"]
    properties.update({
        "yaxis_ticks_angle": yaxis_ticks["tickAngle"],
        "yaxis_ticks_position": yaxis_ticks["tickPosition"],
        "yaxis_ticks_style_mode": yaxis_ticks_styling["currentStylingMode"],
        # Default Y-Axis Tick Label Styling
        "yaxis_ticks_colour_default": yaxis_ticks_styling["customFontStyle"]["fontColour"],
        "yaxis_ticks_size_default": yaxis_ticks_styling["customFontStyle"]["fontSize"],
        "yaxis_ticks_typeface_default": yaxis_ticks_styling["customFontStyle"]["typeface"],
        # Custom Y-Axis Tick Label Styling
        "yaxis_ticks_colour_custom": yaxis_ticks_styling["customFontStyle"]["fontColour"],
        "yaxis_ticks_size_custom": yaxis_ticks_styling["customFontStyle"]["fontSize"],
        "yaxis_ticks_typeface_custom": yaxis_ticks_styling["customFontStyle"]["typeface"],
    })


    # Visual Options - Colour
    visual_options = config_json["visualOptions"]["colour"]
    properties.update({
        "plot_background_colour": visual_options["plotBackgroundColourHex"],
        "plot_margin_colour": visual_options["plotMarginColourHex"],
    })


    return properties


def data_extract(data_json, plot_data_x, plot_data_y) -> dict:
    """
    Extracts the data from json format to a numpy DataFrame.
    :param data_json: The Json file that data is to be extracted from.
    :param name: Name of the data field to be put in a DF.
    :param first_value: If a value in the data json is specified this \
    will be the y axis of the graph.
    :returns: a Pandas Data Frame
    """

    graph_data = {}

    graph_data = extract_data_for_dataframe(graph_data, data_json["month"], plot_data_x, "x")
    graph_data = extract_data_for_dataframe(graph_data, data_json["month"], plot_data_y, "y")

    return graph_data




def extract_data_for_dataframe(graph_data, data_json, search_string, column):
    """
    Extracts values from a nested dictionary based on a given path.

    :param data: The nested dictionary to search.
    :param path: The string path representing the nested structure, separated by dots.
    :return: A list of values found at the path, or an empty list if the path is invalid.
    """
    elements = search_string.split('.')
    current_data = data_json

    for element in elements[:-1]:  # Iterate through the path except the last element
        if isinstance(current_data, dict) and element in current_data:
            current_data = current_data[element]
        elif isinstance(current_data, list) and element.isdigit():
            index = int(element)
            if 0 <= index < len(current_data):
                current_data = current_data[index]
            else:
                return []  # Invalid index
        else:
            return []  # Invalid path

    # Extract the final values based on the last element in the path
    final_element = elements[-1]
    if isinstance(current_data, list):
        my_list = list(dict.fromkeys(
            [item.get(final_element, None)
            for item in current_data if final_element in item]))
        graph_data[column] = my_list
        return graph_data

    return []  # The path does not lead to a list


def pascal_split_name(value:str) -> str:
    """Function for turning PascalCase to normal english"""
    if value is None:
        return None
    if len(value) <= 1:
        return value

    result = []
    i = 0

    while i < len(value):
        if value[i].isupper():
            if i > 0 and value[i - 1].islower():
                result.append(' ')
            elif i > 0 and i + 1 < len(value) and value[i + 1].islower():
                result.append(' ')
            result.append(value[i])
        else:
            result.append(value[i])
        i += 1

    return ''.join(result).strip()


def return_docx(plots,data_json):
    """Takes in a list of plots and returns a docx file with the png's of those plots"""


    doc = Document()
    main_title = doc.add_paragraph()
    main_title_run = main_title.add_run('Plots')
    main_title_run.font.size = Pt(32)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for plot in plots:
        a,b = generate_plot_png(plot, data_json)
        a = io.BytesIO(a)
        plot_title = doc.add_paragraph()
        plot_title_run = plot_title.add_run(b)
        plot_title_run.font.size = Pt(21)
        doc.add_picture(a, width = Inches(6))

    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    return doc_bytes


def delete_folder(folder_path):
    """Delete a folder and all its contents."""
    try:
        # Check if the folder exists
        if os.path.exists(folder_path):
            # Remove the folder and all its contents
            shutil.rmtree(folder_path)
            print(f"Folder '{folder_path}' has been deleted.")
        else:
            print(f"Folder '{folder_path}' does not exist.")
    except PermissionError as e:
        print(f"Permission Denied: {e}")
    except FileNotFoundError as e:
        print(f"File Not Found: {e}")
    except OSError as e:
        print(f"OS Error: {e}")
